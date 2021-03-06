import os
import comtypes.client
import win32com.client as win32
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook
from .Messages.ccs_html_template import ccs_html_template
from .Messages.ces_html_template import ces_html_template

class CertificateMailer():
	"""
		Attributes:
			base_dir: The base directory
			certification_type: Certification type
			certification_year: Certification year
	"""

	def __init__(self, base_dir, certification_type, certification_year):
		"""
		The constructor for the Certificate Class

		Args:
			base_dir(str): Base directory
			certification_type(str): Certification type
			certification_year(str): Certification year
		"""

		self.base_dir = base_dir
		self.certification_type = certification_type
		self.certification_year = certification_year

	def load_excel_sheet(self):
		wb = load_workbook(
			    self.base_dir +
			    '/Lists/' +
			    self.certification_type +
			    '_renewal_list.xlsx')
		sheet = wb.get_sheet_by_name('Pending')

		return sheet

	def create_certificate(self, first_name, last_name):
		"""
		A function that leverages the class' attributes to output a customized
		certificate for each individual.

		Args:
			first_name(str): First name
			last_name(str): Last name
		"""

		full_name = first_name + " " + last_name

		# Loads excel file and certificate attachment
		word = comtypes.client.CreateObject('Word.Application')
		document = Document(
				   self.base_dir +
				   '/Modules/Templates/' +
				   self.certification_type.lower() +
				   '_template_' +
				   self.certification_year +
				   '.docx')

		# Getting fonts ready for the certificate
		obj_styles = document.styles
		obj_charstyle = obj_styles.add_style('Old English Text MT', WD_STYLE_TYPE.CHARACTER)
		obj_font = obj_charstyle.font
		obj_font.size = Pt(48)
		obj_font.name = 'Old English Text MT'

		# Loops through the word document's contents
		for i in document.paragraphs:
			# If the placeholder exists, run this code
			if "{{name}}" in i.text:
				# Sets the text style
				i.text = i.add_run(full_name, style='Old English Text MT').bold = False

				# Save document as .docx - acts as a placeholder
				document.save('placeholder.docx')

				pdf_document = word.Documents.Open(self.base_dir + '\\placeholder.docx')

				# Document is saved as pdf in designated CCS or CES folders
				pdf_document.SaveAs(
						    self.base_dir +
						    "\\Certificates\\" +
						    self.certification_type + 
						    "\\" +
						    full_name +
						    '_' +
						    self.certification_year +
						    ' ' +
						    self.certification_type +
						    " Certificate",
						    FileFormat = 17)

				pdf_document.Close()

				i.text="{{name}}"

				# Prints each individual's name
				print(full_name + "'s certificate has been created")

	def create_and_send_email(
						    self,
						    first_name,
						    last_name,
						    email,
						    ncbfaa_id,
						    renewal_date):
		"""
		A function that creates and sends renewal emails based on the data that passed.

		Args:
			first_name(str): First name
			last_name(str): Last name
			email(str): Email
			ncbfaa_id(int || str): NCBFAA ID
			renewal_date(date): Renewal Date
		"""
		# Converts renewal date into readable string
		renewal_date= renewal_date.strftime('%m/%d/%Y')

		# Full name
		full_name = first_name + " " + last_name

		# Converts ID to string
		ncbfaa_id= str(ncbfaa_id)

		# Opens outlook
		outlook= win32.Dispatch('outlook.application')

		# Gets email ready
		mail= outlook.CreateItem(0)
		mail.To= email

		mail.Subject= 'NCBFAA - ' + self.certification_year + ' ' + self.certification_type + ' Renewal'

		if self.certification_type == 'CCS':
			mail.HTMLBody = ccs_html_template(self.certification_year, first_name, ncbfaa_id, renewal_date)
		elif self.certification_type == 'CES':
			mail.HTMLBody = ces_html_template(self.certification_year, first_name, ncbfaa_id, renewal_date)
		else:
			print('Invalid Certification Type')

		attachment = self.base_dir + "\\Certificates\\" + self.certification_type + "\\" + full_name + '_' + self.certification_year + ' ' + self.certification_type + " Certificate.pdf"
		
		mail.Attachments.Add(Source=attachment)

		mail.Send()

	def mailer(self):
		"""
			A function that loops through each row of a given workbook, creates certificates, and sends
			them out attached to a renewal email.
		"""
		sheet = self.load_excel_sheet()

		# Setting to True, so excel skips the first line (excel header)
		first_line = True

		for row in sheet:
			# Skip first line, switch to false
			if first_line:
				first_line = False
				continue
			"""
				row[2] = first_name,
				row[1] = last_name,
				row[3] = email,
				row[4] = renewal_date,
				row[0] = ncbfaa_id
			"""
			self.create_certificate(row[2].value, row[1].value)
			self.create_and_send_email(
							    row[2].value,
							    row[1].value,
							    row[3].value,
							    row[0].value,
							    row[4].value)

			list_values = []
			full_name = row[2].value + " " + row[1].value
			list_values.append(full_name + ": Completed")

		return list_values