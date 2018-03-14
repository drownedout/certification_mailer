import os
import settings
import comtypes.client
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

class Certificate():

	"""
		Attributes:
			base_dir: The base directory
			certification_type: Certification type
			certification_year: Certification year
			first_name: Individual's first name
			last_name: Individual's last name
	"""

	def __init__(self, base_dir, certification_type, certification_year, first_name, last_name):
		
		"""
		The constructor for the Certificate Class

		Args:
			base_dir(str): Base directory
			certification_type(str): Certification type
			certification_year(str): Certification year
			first_name(str): First name
			last_name(str): Last name 
		"""

		self.base_dir = base_dir
		self.certification_type = certification_type
		self.certification_year = certification_year
		self.full_name = first_name + ' ' + last_name

	def create_certificate(self):

		"""
		A function that leverages the class' attributes to output a customized
		certificate for each individual.
		"""

		# Loads excel file and certificate attachment
		word = comtypes.client.CreateObject('Word.Application')
		document = Document(self.base_dir + '/Templates/' + self.certification_type.lower() + '_template_' + self.certification_year + '.docx')
		
		# Getting fonts ready for the certificate
		obj_styles = document.styles
		obj_charstyle = obj_styles.add_style('Old English Text MT', WD_STYLE_TYPE.CHARACTER)
		obj_font = obj_charstyle.font
		obj_font.size = Pt(48)
		obj_font.name = 'Old English Text MT'

		# Loops through the word document's contents
		for i in document.paragraphs:
			# If the placeholder exists, run this code
			if "{{name}}" in i.text:
				# Sets the text style
				i.text = i.add_run(self.full_name, style = 'Old English Text MT').bold = False

				# Save document as .docx
				document.save('placeholder.docx')

				pdf_document = word.Documents.Open(self.base_dir + 'placeholder.docx'))

				# Document is saved as pdf in designated CCS or CES folders
				pdf_document.SaveAs(self.base_dir + "\\Certificates\\" + self.certification_type + "\\" + self.full_name, FileFormat=17)
				pdf_document.Close()

				i.text = "{{name}}"

				# Prints each individual's name
				print(self.full_name + "'s certificate has been created")